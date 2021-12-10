import docx
import requests
from bs4 import BeautifulSoup

doc = docx.Document()

baseURL = 'https://www.tutorialspoint.com'

url = 'https://www.tutorialspoint.com/artificial_intelligence_with_python/index.htm'

while(url):

    page = requests.get(url)

    soup= BeautifulSoup(page.content,'html.parser')

    contain = soup.find('div', class_= 'mui-container-fluid')

    result = contain.find('div',class_='mui-container')

    p_tags = result.find_all(['p','pre','h2','h3'])

    nxt_btn = result.find('div',class_='nxt-btn')

    nxt_btn_anchor = nxt_btn.find('a')

    nxt_btn_href = nxt_btn_anchor['href']

    if (nxt_btn_href.split('/')[2]=='index.htm'):
        
        break

    for p_tag in p_tags:
        if(p_tag.name == 'h2' or p_tag.name == 'h3'):
             para = doc.add_paragraph()
             bold_para = para.add_run(p_tag.text)
             # Setting bold to true
             bold_para.bold = True
        
        else:
            doc.add_paragraph(p_tag.text)

    url = baseURL + nxt_btn_href


doc.save('ScrapedData.docx')    